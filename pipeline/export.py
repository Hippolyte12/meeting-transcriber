from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor

SPEAKER_COLORS = [
    RGBColor(0x1E, 0x90, 0xFF),  # bleu
    RGBColor(0xFF, 0x45, 0x00),  # rouge-orangé
    RGBColor(0x32, 0xCD, 0x32),  # vert
    RGBColor(0xFF, 0xA5, 0x00),  # orange
    RGBColor(0x94, 0x00, 0xD3),  # violet
]


def _format_time(seconds):
    total = int(seconds)
    h = total // 3600
    m = (total % 3600) // 60
    s = total % 60
    return f"{h:02d}:{m:02d}:{s:02d}"


def _speaker_stats(segments):
    stats = {}
    for seg in segments:
        if seg.speaker not in stats:
            stats[seg.speaker] = [0.0, 0]
        stats[seg.speaker][0] += seg.end - seg.start
        stats[seg.speaker][1] += 1
    return {k: tuple(v) for k, v in sorted(stats.items(), key=lambda x: -x[1][0])}


def _build_lines(segments, with_timestamps=True):
    result = []
    for seg in segments:
        if with_timestamps:
            line = f"[{_format_time(seg.start)}] {seg.speaker} : {seg.text}"
        else:
            line = f"{seg.speaker} : {seg.text}"
        result.append(line)
    return result


def export_txt(segments, output_path, with_timestamps=True, title="Réunion"):
    lines = _build_lines(segments, with_timestamps)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(f"{title}\n")
        f.write("=" * len(title) + "\n\n")
        for line in lines:
            f.write(line + "\n")


def export_markdown(segments, output_path, with_timestamps=True, title="Réunion"):
    stats = _speaker_stats(segments)
    lines = _build_lines(segments, with_timestamps)

    with open(output_path, "w", encoding="utf-8") as f:
        f.write(f"# {title}\n\n")
        f.write("## Statistiques\n\n")
        f.write("| Locuteur | Temps de parole | Interventions |\n")
        f.write("|----------|----------------|---------------|\n")
        for speaker, (duration, count) in stats.items():
            f.write(f"| {speaker} | {_format_time(duration)} | {count} |\n")
        f.write("\n## Transcript\n\n")
        for line in lines:
            f.write(line + "\n")


def export_docx(segments, output_path, with_timestamps=True, title="Réunion"):
    doc = Document()
    doc.add_heading(title, level=1)

    # Mapping locuteur → couleur
    speakers = list(dict.fromkeys(seg.speaker for seg in segments))
    color_map = {s: SPEAKER_COLORS[i % len(SPEAKER_COLORS)] for i, s in enumerate(speakers)}

    for seg in segments:
        para = doc.add_paragraph()

        if with_timestamps:
            run_time = para.add_run(f"[{_format_time(seg.start)}] ")
            run_time.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            run_time.font.size = Pt(9)

        run_speaker = para.add_run(f"{seg.speaker} : ")
        run_speaker.bold = True
        run_speaker.font.color.rgb = color_map[seg.speaker]

        para.add_run(seg.text)

    doc.save(output_path)